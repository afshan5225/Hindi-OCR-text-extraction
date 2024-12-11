import streamlit as st
import json
import re
from pdf2image import convert_from_path
import numpy as np
import pytesseract
import cv2
from langchain_groq import ChatGroq
from docx import Document
import os

# Function to clean and process the OCR text
def clean_text(text):
    text = re.sub(r"\n", " ", text)  # Replace newline with space
    text = re.sub(r"[|\[\]]", "", text)  # Remove stray characters like |, [, ]
    text = re.sub(r"\s+", " ", text).strip()  # Normalize whitespace
    return text

# Function to extract text from the image using Tesseract OCR
def extract_text(image):
    custom_config = r'--oem 3 --psm 6 -l hin'  # Hindi language OCR configuration
    text = pytesseract.image_to_string(image, config=custom_config)
    return clean_text(text)

# Function to query the LLM with cleaned text
def query_llm(cleaned_text): 
    llm = ChatGroq(
        temperature=0,
        groq_api_key='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',
        model_name="llama-3.1-70b-versatile"
    )

    prompt = f"""
    The following text has been extracted from a document:
    
    {cleaned_text}
    
    Please structure it into the following format:
    - Headers: Any section titles or headers.
    - Paragraphs: Normal text grouped into meaningful paragraphs.
    - Tables: Tabular data (rows and columns) presented as JSON.
    - Notes: Any additional observations.
    \n\n\nfor example\n\n\n: the extracted text will be in Hindi, and that will be in the form: 
    "विधान सभा निर्वाचन क्षेत्र, बिहार की निर्वाचक नामावली : भाग संख्या: प्रभाग की संख्या व नाम: 2 निर्वाचक का नाम : सुबीदा देवी निर्वाचक का नाम ; उमा देवी 55.0. पति का नाम: सदन प्रसाद पति का नाम: ईश्वार प्रसाद पिता का नाम: जग्गु प्रसाद गृह संख्या फोटो उपलब्ध गृह संख्या: 2 फोटो उपलब्ध गृह संख्या: 2 फोटो उपलब्ध उम्र : 68 लिंग : महिला उम्र : 79 लिंग : महिला उम्र ; 73 लिंग : पुरुष"
    you will have to arrange them like in this format: 
    
    निर्वांचक का नाम : सुबीदा देवी
    पति का नाम: सदन प्रसाद
    गृह संख्या: 2
    उम्र : 68
    लिंग : महिला 
    
    Return the result as a JSON object.
    """

    response = llm.invoke(prompt)
    return response.content

# Function to chunk text and send to LLM
def chunk_text(text, chunk_size=6000):
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
    return chunks

# Function to extract pages from PDF, apply inversion, OCR, and query the LLM
def extract_pages_and_invert(pdf_path, start_page, end_page):
    pages = convert_from_path(pdf_path, dpi=300)
    combined_text = ""
    
    for page_num in range(start_page - 1, end_page):  # 0-indexed
        img = pages[page_num]
        
        img = np.array(img)
        inverted_image = cv2.bitwise_not(img)
        
        inverted_text = extract_text(inverted_image)
        
        combined_text += inverted_text + " "
    
    cleaned_text = clean_text(combined_text)
    
    # Chunk the cleaned text if it's too large for the model
    text_chunks = chunk_text(cleaned_text)
    
    data = {}

    for idx, chunk in enumerate(text_chunks):
        structured_response = query_llm(chunk)
        data[f"chunk_{idx + 1}"] = structured_response
    
    extracted_json = {}

    # Regex pattern to extract the JSON block
    pattern = r'(\{.*\})'

    for chunk_key, chunk_data in data.items():
        match = re.search(pattern, chunk_data, re.DOTALL)

        if match:
            cleaned = match.group(1)
            extracted_json[chunk_key] = json.loads(cleaned)
        else:
            extracted_json[chunk_key] = None

    return extracted_json

# Streamlit UI setup
st.title("OCR and Data Extraction Application")
uploaded_file = st.file_uploader("Upload a PDF", type="pdf")

if uploaded_file is not None:
    start_page = st.number_input("Start Page", min_value=1, value=1)
    end_page = st.number_input("End Page", min_value=start_page, value=start_page)

    if st.button("Extract Data"):
        # Save uploaded file to disk
        pdf_path = f"uploaded_pdf.pdf"
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.read())

        # Extract and process the PDF
        data = extract_pages_and_invert(pdf_path, start_page, end_page)

        # Prepare the DOCX file
        docx_file_path = os.path.join(os.environ["USERPROFILE"], "Desktop", "voter_report.docx")
        document = Document()
        
        # Add Header
        header = "विधान सभा निर्वाचन क्षेत्र, बिहार की निर्वाचक नामावली : 7-अस्थावॉँ भाग संख्या : : प्रभाग की संख्या व नाम : -जक्की"
        doc_paragraph = document.add_paragraph()
        doc_paragraph.add_run(header).bold = True
        doc_paragraph.style = 'Title'

        # Add a Table
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Define header row for the table
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "निर्वाचक का नाम"
        hdr_cells[1].text = "पति का नाम / पिता का नाम"
        hdr_cells[2].text = "गृह संख्या"
        hdr_cells[3].text = "उम्र"
        hdr_cells[4].text = "लिंग"

        # Loop through the JSON table entries to populate table rows
        for chunk_key, chunk_value in data.items():
            if chunk_value:
                for entry in chunk_value.get('Tables', []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = entry.get("निर्वाचक का नाम", " ")
                    row_cells[1].text = entry.get("पति का नाम", entry.get("पिता का नाम", ""))
                    row_cells[2].text = entry.get("गृह संख्या", "")
                    row_cells[3].text = entry.get("उम्र", "")
                    row_cells[4].text = entry.get("लिंग", "")

        # Add Notes Section
        

        # Save the DOCX file
        document.save(docx_file_path)

        st.success(f"DOCX file has been created successfully at: {docx_file_path}")
        st.download_button("Download DOCX", docx_file_path, file_name="voter_report.docx")
